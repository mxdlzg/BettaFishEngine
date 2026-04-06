# -*- coding: utf-8 -*-
"""
LangGraph版本的Deep Search Agent

使用LangGraph StateGraph重构QueryEngine，提供标准的图执行流程。
"""

import os
from typing import TypedDict, List, Dict, Any, Optional, Annotated
from datetime import datetime

from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from loguru import logger

from .llms import LLMClient
from .nodes import (
    ReportStructureNode,
    FirstSearchNode, 
    ReflectionNode,
    FirstSummaryNode,
    ReflectionSummaryNode,
    ReportFormattingNode
)
from .tools import DuckDuckGoNewsAgency
from .utils import Settings, format_search_results_for_prompt


# ============== LangGraph State 定义 ==============

class SearchResult(TypedDict):
    """搜索结果"""
    query: str
    url: str
    title: str
    content: str
    score: Optional[float]
    published_date: Optional[str]


class ParagraphState(TypedDict):
    """段落状态"""
    title: str
    content: str
    search_history: List[SearchResult]
    latest_summary: str
    reflection_count: int
    is_completed: bool


class DeepSearchState(TypedDict):
    """LangGraph状态定义"""
    # 输入
    query: str
    
    # 配置
    max_reflections: int
    export_formats: List[str]
    
    # 中间状态
    report_title: str
    paragraphs: List[ParagraphState]
    current_paragraph_index: int
    current_reflection_index: int
    
    # 搜索相关
    current_search_query: str
    current_search_tool: str
    current_search_results: List[SearchResult]
    
    # 输出
    final_report: str
    output_dir: str
    saved_files: Dict[str, str]
    
    # 状态标志
    is_completed: bool
    error: Optional[str]


# ============== 节点函数 ==============

def create_deep_search_graph(config: Optional[Settings] = None) -> StateGraph:
    """
    创建Deep Search LangGraph
    
    Args:
        config: 配置对象
        
    Returns:
        编译后的StateGraph
    """
    from .utils.config import settings
    config = config or settings
    
    # 初始化共享资源
    llm_client = _initialize_llm(config)
    search_agency = DuckDuckGoNewsAgency()
    
    # 初始化所有节点
    nodes = {
        'first_search': FirstSearchNode(llm_client),
        'reflection': ReflectionNode(llm_client),
        'first_summary': FirstSummaryNode(llm_client),
        'reflection_summary': ReflectionSummaryNode(llm_client),
        'formatting': ReportFormattingNode(llm_client)
    }
    
    # ============== 节点实现 ==============
    
    def intake_node(state: DeepSearchState) -> DeepSearchState:
        """入口节点：初始化状态"""
        logger.info(f"\n{'='*60}")
        logger.info(f"开始深度研究: {state['query']}")
        logger.info(f"{'='*60}")
        
        return {
            **state,
            "paragraphs": [],
            "current_paragraph_index": 0,
            "current_reflection_index": 0,
            "is_completed": False,
            "error": None
        }
    
    def generate_structure_node(state: DeepSearchState) -> DeepSearchState:
        """生成报告结构节点"""
        logger.info(f"\n[步骤 1] 生成报告结构...")
        
        try:
            structure_node = ReportStructureNode(llm_client, state['query'])
            result = structure_node.run()
            
            paragraphs = []
            for i, item in enumerate(result):
                paragraphs.append({
                    "title": item.get("title", f"段落{i+1}"),
                    "content": item.get("content", ""),
                    "search_history": [],
                    "latest_summary": "",
                    "reflection_count": 0,
                    "is_completed": False
                })
            
            logger.info(f"报告结构已生成，共 {len(paragraphs)} 个段落")
            for i, p in enumerate(paragraphs, 1):
                logger.info(f"  {i}. {p['title']}")
            
            return {
                **state,
                "paragraphs": paragraphs,
                "report_title": f"【深度调查】{state['query']}全面新闻分析报告"
            }
        except Exception as e:
            logger.error(f"生成报告结构失败: {e}")
            return {**state, "error": str(e)}
    
    def initial_search_node(state: DeepSearchState) -> DeepSearchState:
        """初始搜索节点：为当前段落生成搜索查询"""
        idx = state['current_paragraph_index']
        paragraph = state['paragraphs'][idx]
        
        logger.info(f"\n[步骤 2.{idx+1}] 处理段落: {paragraph['title']}")
        logger.info("-" * 50)
        logger.info("  - 生成搜索查询...")
        
        search_input = {
            "title": paragraph['title'],
            "content": paragraph['content']
        }
        
        search_output = nodes['first_search'].run(search_input)
        
        logger.info(f"  - 搜索查询: {search_output['search_query']}")
        logger.info(f"  - 选择的工具: {search_output.get('search_tool', 'basic_search_news')}")
        
        return {
            **state,
            "current_search_query": search_output['search_query'],
            "current_search_tool": search_output.get('search_tool', 'basic_search_news')
        }
    
    def execute_search_node(state: DeepSearchState) -> DeepSearchState:
        """执行搜索节点"""
        logger.info("  - 执行网络搜索...")
        
        tool_name = state['current_search_tool']
        query = state['current_search_query']
        
        # 执行搜索
        search_response = _execute_search(search_agency, tool_name, query)
        
        # 转换结果
        search_results = []
        if search_response and search_response.results:
            for result in search_response.results[:10]:
                search_results.append({
                    'query': query,
                    'title': result.title,
                    'url': result.url,
                    'content': result.content,
                    'score': result.score,
                    'published_date': result.published_date
                })
        
        logger.info(f"  - 找到 {len(search_results)} 个搜索结果")
        
        return {
            **state,
            "current_search_results": search_results
        }
    
    def initial_summary_node(state: DeepSearchState) -> DeepSearchState:
        """初始总结节点"""
        logger.info("  - 生成初始总结...")
        
        idx = state['current_paragraph_index']
        paragraph = state['paragraphs'][idx]
        
        summary_input = {
            "title": paragraph['title'],
            "content": paragraph['content'],
            "search_query": state['current_search_query'],
            "search_results": format_search_results_for_prompt(
                state['current_search_results'], 
                config.SEARCH_CONTENT_MAX_LENGTH
            )
        }
        
        # run() 返回字符串（段落内容）
        summary_content = nodes['first_summary'].run(summary_input)
        
        # 更新段落状态
        paragraphs = state['paragraphs'].copy()
        paragraphs[idx] = {
            **paragraphs[idx],
            "latest_summary": summary_content if isinstance(summary_content, str) else str(summary_content),
            "search_history": paragraphs[idx]['search_history'] + state['current_search_results']
        }
        
        logger.info("  - 初始总结完成")
        
        return {
            **state,
            "paragraphs": paragraphs
        }
    
    def reflection_search_node(state: DeepSearchState) -> DeepSearchState:
        """反思搜索节点"""
        idx = state['current_paragraph_index']
        ref_idx = state['current_reflection_index']
        paragraph = state['paragraphs'][idx]
        
        logger.info(f"  - 反思 {ref_idx + 1}/{state['max_reflections']}...")
        
        reflection_input = {
            "title": paragraph['title'],
            "content": paragraph['content'],
            "paragraph_latest_state": paragraph['latest_summary']
        }
        
        reflection_output = nodes['reflection'].run(reflection_input)
        
        logger.info(f"    反思查询: {reflection_output['search_query']}")
        
        return {
            **state,
            "current_search_query": reflection_output['search_query'],
            "current_search_tool": reflection_output.get('search_tool', 'basic_search_news')
        }
    
    def reflection_summary_node(state: DeepSearchState) -> DeepSearchState:
        """反思总结节点"""
        idx = state['current_paragraph_index']
        paragraph = state['paragraphs'][idx]
        
        summary_input = {
            "title": paragraph['title'],
            "content": paragraph['content'],
            "search_query": state['current_search_query'],
            "search_results": format_search_results_for_prompt(
                state['current_search_results'],
                config.SEARCH_CONTENT_MAX_LENGTH
            ),
            "paragraph_latest_state": paragraph['latest_summary']
        }
        
        # run() 返回字符串（段落内容）
        summary_content = nodes['reflection_summary'].run(summary_input)
        
        # 更新段落状态
        paragraphs = state['paragraphs'].copy()
        paragraphs[idx] = {
            **paragraphs[idx],
            "latest_summary": summary_content if isinstance(summary_content, str) else paragraph['latest_summary'],
            "search_history": paragraphs[idx]['search_history'] + state['current_search_results'],
            "reflection_count": state['current_reflection_index'] + 1
        }
        
        logger.info(f"    反思 {state['current_reflection_index'] + 1} 完成")
        
        return {
            **state,
            "paragraphs": paragraphs,
            "current_reflection_index": state['current_reflection_index'] + 1
        }
    
    def check_reflection_complete(state: DeepSearchState) -> str:
        """检查反思是否完成"""
        if state['current_reflection_index'] >= state['max_reflections']:
            return "reflection_done"
        return "continue_reflection"
    
    def mark_paragraph_complete_node(state: DeepSearchState) -> DeepSearchState:
        """标记当前段落完成"""
        idx = state['current_paragraph_index']
        
        paragraphs = state['paragraphs'].copy()
        paragraphs[idx] = {
            **paragraphs[idx],
            "is_completed": True
        }
        
        progress = (idx + 1) / len(paragraphs) * 100
        logger.info(f"段落处理完成 ({progress:.1f}%)")
        
        return {
            **state,
            "paragraphs": paragraphs,
            "current_paragraph_index": idx + 1,
            "current_reflection_index": 0
        }
    
    def check_all_paragraphs_done(state: DeepSearchState) -> str:
        """检查所有段落是否完成"""
        if state['current_paragraph_index'] >= len(state['paragraphs']):
            return "all_done"
        return "next_paragraph"
    
    def generate_report_node(state: DeepSearchState) -> DeepSearchState:
        """生成最终报告节点"""
        logger.info(f"\n[步骤 3] 生成最终报告...")
        
        report_data = []
        for paragraph in state['paragraphs']:
            report_data.append({
                "title": paragraph['title'],
                "paragraph_latest_state": paragraph['latest_summary']
            })
        
        try:
            final_report = nodes['formatting'].run(report_data)
        except Exception as e:
            logger.error(f"LLM格式化失败，使用备用方法: {e}")
            final_report = nodes['formatting'].format_report_manually(
                report_data, state['report_title']
            )
        
        logger.info("最终报告生成完成")
        
        return {
            **state,
            "final_report": final_report
        }
    
    def save_report_node(state: DeepSearchState) -> DeepSearchState:
        """保存报告节点"""
        import uuid
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        task_id = uuid.uuid4().hex[:8]
        query_safe = "".join(c for c in state['query'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        query_safe = query_safe.replace(' ', '_')[:30]
        
        task_folder = os.path.join(config.OUTPUT_DIR, f"task_{query_safe}_{timestamp}_{task_id}")
        os.makedirs(task_folder, exist_ok=True)
        
        saved_files = {}
        export_formats = state.get('export_formats', ['md', 'html'])
        
        # 保存Markdown
        if 'md' in export_formats:
            md_path = os.path.join(task_folder, "report.md")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(state['final_report'])
            saved_files['md'] = md_path
            logger.info(f"Markdown报告已保存到: {md_path}")
        
        # 保存HTML
        if 'html' in export_formats:
            try:
                import markdown
                html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{state['report_title']}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
               max-width: 900px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1, h2, h3, h4 {{ color: #333; }}
        a {{ color: #0066cc; }}
        pre {{ background: #f5f5f5; padding: 15px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #ddd; margin-left: 0; padding-left: 20px; color: #666; }}
    </style>
</head>
<body>
{markdown.markdown(state['final_report'], extensions=['tables', 'fenced_code', 'toc'])}
</body>
</html>"""
                html_path = os.path.join(task_folder, "report.html")
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                saved_files['html'] = html_path
                logger.info(f"HTML报告已保存到: {html_path}")
            except Exception as e:
                logger.error(f"HTML导出失败: {e}")
        
        # 保存DOCX
        if 'docx' in export_formats:
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document()
                
                for line in state['final_report'].split('\n'):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('#### '):
                        doc.add_heading(line[5:], level=4)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line)
                
                docx_path = os.path.join(task_folder, "report.docx")
                doc.save(docx_path)
                saved_files['docx'] = docx_path
                logger.info(f"Word报告已保存到: {docx_path}")
            except Exception as e:
                logger.error(f"DOCX导出失败: {e}")
        
        # 保存manifest
        import json
        manifest = {
            "task_id": task_id,
            "query": state['query'],
            "timestamp": timestamp,
            "files": saved_files
        }
        manifest_path = os.path.join(task_folder, "manifest.json")
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        
        logger.info(f"任务报告已保存到文件夹: {task_folder}")
        
        return {
            **state,
            "output_dir": task_folder,
            "saved_files": saved_files,
            "is_completed": True
        }
    
    def finalize_node(state: DeepSearchState) -> DeepSearchState:
        """最终节点"""
        logger.info(f"\n{'='*60}")
        logger.info("深度研究完成！")
        logger.info(f"{'='*60}")
        return state
    
    # ============== 构建图 ==============
    
    builder = StateGraph(DeepSearchState)
    
    # 添加节点
    builder.add_node("intake", intake_node)
    builder.add_node("generate_structure", generate_structure_node)
    builder.add_node("initial_search", initial_search_node)
    builder.add_node("execute_search", execute_search_node)
    builder.add_node("initial_summary", initial_summary_node)
    builder.add_node("reflection_search", reflection_search_node)
    builder.add_node("reflection_execute_search", execute_search_node)
    builder.add_node("reflection_summary", reflection_summary_node)
    builder.add_node("mark_paragraph_complete", mark_paragraph_complete_node)
    builder.add_node("generate_report", generate_report_node)
    builder.add_node("save_report", save_report_node)
    builder.add_node("finalize", finalize_node)
    
    # 定义边
    builder.add_edge(START, "intake")
    builder.add_edge("intake", "generate_structure")
    builder.add_edge("generate_structure", "initial_search")
    builder.add_edge("initial_search", "execute_search")
    builder.add_edge("execute_search", "initial_summary")
    builder.add_edge("initial_summary", "reflection_search")
    builder.add_edge("reflection_search", "reflection_execute_search")
    builder.add_edge("reflection_execute_search", "reflection_summary")
    
    # 条件边：检查反思是否完成
    builder.add_conditional_edges(
        "reflection_summary",
        check_reflection_complete,
        {
            "continue_reflection": "reflection_search",
            "reflection_done": "mark_paragraph_complete"
        }
    )
    
    # 条件边：检查所有段落是否完成
    builder.add_conditional_edges(
        "mark_paragraph_complete",
        check_all_paragraphs_done,
        {
            "next_paragraph": "initial_search",
            "all_done": "generate_report"
        }
    )
    
    builder.add_edge("generate_report", "save_report")
    builder.add_edge("save_report", "finalize")
    builder.add_edge("finalize", END)
    
    return builder.compile()


# ============== 辅助函数 ==============

def _initialize_llm(config: Settings) -> LLMClient:
    """初始化LLM客户端"""
    return LLMClient(
        api_key=config.QUERY_ENGINE_API_KEY,
        base_url=config.QUERY_ENGINE_BASE_URL,
        model_name=config.QUERY_ENGINE_MODEL_NAME
    )


def _execute_search(search_agency: DuckDuckGoNewsAgency, tool_name: str, query: str, **kwargs):
    """执行搜索"""
    tool_map = {
        "basic_search_news": search_agency.basic_search_news,
        "search_news_last_24_hours": search_agency.search_news_last_24_hours,
        "search_news_last_week": search_agency.search_news_last_week,
        "search_news_by_date": search_agency.search_news_by_date,
        "deep_search_news": search_agency.deep_search_news,
    }
    
    search_func = tool_map.get(tool_name, search_agency.basic_search_news)
    
    try:
        return search_func(query, **kwargs)
    except Exception as e:
        logger.warning(f"搜索工具 {tool_name} 执行失败: {e}")
        return search_agency.basic_search_news(query)


# ============== 便捷调用函数 ==============

def run_deep_search(
    query: str,
    max_reflections: int = 3,
    export_formats: List[str] = None,
    config: Optional[Settings] = None
) -> Dict[str, Any]:
    """
    运行深度搜索（LangGraph版本）
    
    Args:
        query: 研究查询
        max_reflections: 每个段落的最大反思次数
        export_formats: 导出格式列表 ['md', 'html', 'docx', 'pdf']
        config: 配置对象
        
    Returns:
        包含报告和元数据的字典
    """
    if export_formats is None:
        export_formats = ['md', 'html', 'docx']
    
    # 创建图
    graph = create_deep_search_graph(config)
    
    # 初始状态
    initial_state: DeepSearchState = {
        "query": query,
        "max_reflections": max_reflections,
        "export_formats": export_formats,
        "report_title": "",
        "paragraphs": [],
        "current_paragraph_index": 0,
        "current_reflection_index": 0,
        "current_search_query": "",
        "current_search_tool": "",
        "current_search_results": [],
        "final_report": "",
        "output_dir": "",
        "saved_files": {},
        "is_completed": False,
        "error": None
    }
    
    # 执行图（增加递归限制以支持多段落多反思）
    # 每个段落需要约 6 个节点调用（search, execute, summary, reflection_search, reflection_execute, reflection_summary）
    # 加上反思循环，每个段落约需 6 + (max_reflections * 3) 次调用
    # 5段落 * (6 + 3*3) = 75，再加上结构生成和报告生成，约需100次
    recursion_limit = max(100, 20 * max_reflections * 10)
    
    final_state = graph.invoke(initial_state, {"recursion_limit": recursion_limit})
    
    return {
        "report": final_state["final_report"],
        "output_dir": final_state["output_dir"],
        "saved_files": final_state["saved_files"],
        "is_completed": final_state["is_completed"],
        "error": final_state.get("error")
    }


# 导出
__all__ = [
    "DeepSearchState",
    "create_deep_search_graph",
    "run_deep_search"
]
