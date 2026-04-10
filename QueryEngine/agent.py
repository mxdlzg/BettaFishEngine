"""
Deep Search Agent主类
整合所有模块，实现完整的深度搜索流程
"""

import json
import os
import re
from datetime import datetime
from typing import Optional, Dict, Any, List

from .llms import LLMClient
from .nodes import (
    ReportStructureNode,
    FirstSearchNode, 
    ReflectionNode,
    FirstSummaryNode,
    ReflectionSummaryNode,
    ReportFormattingNode
)
from .state import State
from .tools import TavilyNewsAgency, TavilyResponse
from .utils import Settings, format_search_results_for_prompt
from loguru import logger

class DeepSearchAgent:
    """Deep Search Agent主类"""
    
    def __init__(self, config: Optional[Settings] = None):
        """
        初始化Deep Search Agent
        
        Args:
            config: 配置对象，如果不提供则自动加载
        """
        # 加载配置
        from .utils.config import settings
        self.config = config or settings
        
        # 初始化LLM客户端
        self.llm_client = self._initialize_llm()
        
        # 初始化搜索工具集 - 默认使用Tavily
        self.search_agency = TavilyNewsAgency(api_key=self.config.TAVILY_API_KEY)
        self.search_tool_name = "TavilyNewsAgency"
        logger.info(f"使用搜索工具: {self.search_tool_name}")
        
        # 初始化节点
        self._initialize_nodes()
        
        # 状态
        self.state = State()
        
        # 确保输出目录存在
        os.makedirs(self.config.OUTPUT_DIR, exist_ok=True)
        
        logger.info(f"Query Agent已初始化")
        logger.info(f"使用LLM: {self.llm_client.get_model_info()}")
        logger.info(f"搜索工具集: {self.search_tool_name} (支持6种搜索工具)")
    
    def _initialize_llm(self) -> LLMClient:
        """初始化LLM客户端"""
        return LLMClient(
            api_key=self.config.QUERY_ENGINE_API_KEY,
            model_name=self.config.QUERY_ENGINE_MODEL_NAME,
            base_url=self.config.QUERY_ENGINE_BASE_URL,
        )
    
    def _initialize_nodes(self):
        """初始化处理节点"""
        self.first_search_node = FirstSearchNode(self.llm_client)
        self.reflection_node = ReflectionNode(self.llm_client)
        self.first_summary_node = FirstSummaryNode(self.llm_client)
        self.reflection_summary_node = ReflectionSummaryNode(self.llm_client)
        self.report_formatting_node = ReportFormattingNode(self.llm_client)
    
    def _validate_date_format(self, date_str: str) -> bool:
        """
        验证日期格式是否为YYYY-MM-DD
        
        Args:
            date_str: 日期字符串
            
        Returns:
            是否为有效格式
        """
        if not date_str:
            return False
        
        # 检查格式
        pattern = r'^\d{4}-\d{2}-\d{2}$'
        if not re.match(pattern, date_str):
            return False
        
        # 检查日期是否有效
        try:
            datetime.strptime(date_str, '%Y-%m-%d')
            return True
        except ValueError:
            return False
    
    def execute_search_tool(self, tool_name: str, query: str, **kwargs) -> TavilyResponse:
        """
        执行指定的搜索工具
        
        Args:
            tool_name: 工具名称，可选值：
                - "basic_search_news": 基础新闻搜索（快速、通用）
                - "deep_search_news": 深度新闻分析
                - "search_news_last_24_hours": 24小时内最新新闻
                - "search_news_last_week": 本周新闻
                - "search_images_for_news": 新闻图片搜索
                - "search_news_by_date": 按日期范围搜索新闻
            query: 搜索查询
            **kwargs: 额外参数（如start_date, end_date, max_results）
            
        Returns:
            TavilyResponse对象
        """
        logger.info(f"  → 执行搜索工具: {tool_name}")
        
        if tool_name == "basic_search_news":
            max_results = kwargs.get("max_results", 7)
            return self.search_agency.basic_search_news(query, max_results)
        elif tool_name == "deep_search_news":
            return self.search_agency.deep_search_news(query)
        elif tool_name == "search_news_last_24_hours":
            return self.search_agency.search_news_last_24_hours(query)
        elif tool_name == "search_news_last_week":
            return self.search_agency.search_news_last_week(query)
        elif tool_name == "search_images_for_news":
            return self.search_agency.search_images_for_news(query)
        elif tool_name == "search_news_by_date":
            start_date = kwargs.get("start_date")
            end_date = kwargs.get("end_date")
            if not start_date or not end_date:
                raise ValueError("search_news_by_date工具需要start_date和end_date参数")
            return self.search_agency.search_news_by_date(query, start_date, end_date)
        else:
            logger.warning(f"  ⚠️  未知的搜索工具: {tool_name}，使用默认基础搜索")
            return self.search_agency.basic_search_news(query)
    
    def research(self, query: str, save_report: bool = True, export_formats: List[str] = None) -> str:
        """
        执行深度研究
        
        Args:
            query: 研究查询
            save_report: 是否保存报告到文件
            export_formats: 导出格式列表，可选 ['md', 'html', 'pdf', 'docx']，默认只导出md
            
        Returns:
            最终报告内容
        """
        logger.info(f"\n{'='*60}")
        logger.info(f"开始深度研究: {query}")
        logger.info(f"{'='*60}")
        
        try:
            # Step 1: 生成报告结构
            self._generate_report_structure(query)
            
            # Step 2: 处理每个段落
            self._process_paragraphs()
            
            # Step 3: 生成最终报告
            final_report = self._generate_final_report()
            
            # Step 4: 保存报告
            if save_report:
                saved_files = self._save_report(final_report, export_formats or ['md'])
                if saved_files:
                    logger.info(f"已导出文件: {list(saved_files.keys())}")
            
            logger.info(f"\n{'='*60}")
            logger.info("深度研究完成！")
            logger.info(f"{'='*60}")
            
            return final_report
            
        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            logger.error(f"研究过程中发生错误: {str(e)} \n错误堆栈: {error_traceback}")
            raise e
    
    def _generate_report_structure(self, query: str):
        """生成报告结构"""
        logger.info(f"\n[步骤 1] 生成报告结构...")
        
        # 创建报告结构节点
        report_structure_node = ReportStructureNode(
            self.llm_client,
            query,
            max_paragraphs=int(getattr(self.config, "MAX_PARAGRAPHS", 5) or 5),
        )
        
        # 生成结构并更新状态
        self.state = report_structure_node.mutate_state(state=self.state)
        
        _message = f"报告结构已生成，共 {len(self.state.paragraphs)} 个段落:"
        for i, paragraph in enumerate(self.state.paragraphs, 1):
            _message += f"\n  {i}. {paragraph.title}"
        logger.info(_message)
    
    def _process_paragraphs(self):
        """处理所有段落"""
        from concurrent.futures import ThreadPoolExecutor, as_completed

        total_paragraphs = len(self.state.paragraphs)
        max_paragraphs = max(1, int(getattr(self.config, "MAX_PARAGRAPHS", total_paragraphs) or total_paragraphs))
        total_paragraphs = min(total_paragraphs, max_paragraphs)

        max_workers = max(1, int(getattr(self.config, "MAX_PARAGRAPH_WORKERS", 3) or 3))
        max_workers = min(max_workers, total_paragraphs) if total_paragraphs > 0 else 1

        if total_paragraphs <= 1 or max_workers <= 1:
            for i in range(total_paragraphs):
                self._process_single_paragraph(i, total_paragraphs)
                progress = (i + 1) / total_paragraphs * 100
                logger.info(f"段落处理完成 ({progress:.1f}%)")
            return

        logger.info(f"启用段落并发处理: workers={max_workers}, paragraphs={total_paragraphs}")
        completed = 0
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_index = {
                executor.submit(self._process_single_paragraph, i, total_paragraphs): i
                for i in range(total_paragraphs)
            }
            for future in as_completed(future_to_index):
                future.result()
                completed += 1
                progress = completed / total_paragraphs * 100
                logger.info(f"段落处理完成 ({progress:.1f}%) [{completed}/{total_paragraphs}]")

    def _process_single_paragraph(self, paragraph_index: int, total_paragraphs: int):
        paragraph = self.state.paragraphs[paragraph_index]
        logger.info(f"\n[步骤 2.{paragraph_index+1}] 处理段落: {paragraph.title}")
        logger.info("-" * 50)
        
        # 初始搜索和总结
        self._initial_search_and_summary(paragraph_index)

        # 反思循环
        self._reflection_loop(paragraph_index)

        # 标记段落完成
        paragraph.research.mark_completed()
    
    def _initial_search_and_summary(self, paragraph_index: int):
        """执行初始搜索和总结"""
        paragraph = self.state.paragraphs[paragraph_index]
        
        # 准备搜索输入
        search_input = {
            "title": paragraph.title,
            "content": paragraph.content
        }
        
        # 生成搜索查询和工具选择
        logger.info("  - 生成搜索查询...")
        search_output = self.first_search_node.run(search_input)
        search_query = search_output["search_query"]
        search_tool = search_output.get("search_tool", "basic_search_news")  # 默认工具
        reasoning = search_output["reasoning"]
        
        logger.info(f"  - 搜索查询: {search_query}")
        logger.info(f"  - 选择的工具: {search_tool}")
        logger.info(f"  - 推理: {reasoning}")
        
        # 执行搜索
        logger.info("  - 执行网络搜索...")
        
        # 处理search_news_by_date的特殊参数
        search_kwargs = {}
        if search_tool == "search_news_by_date":
            start_date = search_output.get("start_date")
            end_date = search_output.get("end_date")
            
            if start_date and end_date:
                # 验证日期格式
                if self._validate_date_format(start_date) and self._validate_date_format(end_date):
                    search_kwargs["start_date"] = start_date
                    search_kwargs["end_date"] = end_date
                    logger.info(f"  - 时间范围: {start_date} 到 {end_date}")
                else:
                    logger.info(f"  ⚠️  日期格式错误（应为YYYY-MM-DD），改用基础搜索")
                    logger.info(f"      提供的日期: start_date={start_date}, end_date={end_date}")
                    search_tool = "basic_search_news"
            else:
                logger.info(f"  ⚠️  search_news_by_date工具缺少时间参数，改用基础搜索")
                search_tool = "basic_search_news"
        
        search_response = self.execute_search_tool(search_tool, search_query, **search_kwargs)
        
        # 转换为兼容格式
        search_results = []
        if search_response and search_response.results:
            # Use config-driven cap so gateway max_search_results is effective.
            configured_cap = int(getattr(self.config, "MAX_SEARCH_RESULTS", 10) or 10)
            configured_cap = max(1, configured_cap)
            max_results = min(len(search_response.results), configured_cap)
            for result in search_response.results[:max_results]:
                search_results.append({
                    'title': result.title,
                    'url': result.url,
                    'content': result.content,
                    'score': result.score,
                    'raw_content': result.raw_content,
                    'published_date': result.published_date  # 新增字段
                })
        
        if search_results:
            _message = f"  - 找到 {len(search_results)} 个搜索结果"
            for j, result in enumerate(search_results, 1):
                date_info = f" (发布于: {result.get('published_date', 'N/A')})" if result.get('published_date') else ""
                _message += f"\n    {j}. {result['title'][:50]}...{date_info}"
            logger.info(_message)
        else:
            logger.info("  - 未找到搜索结果")
        # 更新状态中的搜索历史
        paragraph.research.add_search_results(search_query, search_results)
        
        # 生成初始总结
        logger.info("  - 生成初始总结...")
        summary_input = {
            "title": paragraph.title,
            "content": paragraph.content,
            "search_query": search_query,
            "search_results": format_search_results_for_prompt(
                search_results, self.config.SEARCH_CONTENT_MAX_LENGTH
            )
        }
        
        # 更新状态
        self.state = self.first_summary_node.mutate_state(
            summary_input, self.state, paragraph_index
        )
        
        logger.info("  - 初始总结完成")
    
    def _reflection_loop(self, paragraph_index: int):
        """执行反思循环"""
        paragraph = self.state.paragraphs[paragraph_index]
        
        for reflection_i in range(self.config.MAX_REFLECTIONS):
            logger.info(f"  - 反思 {reflection_i + 1}/{self.config.MAX_REFLECTIONS}...")
            
            # 准备反思输入
            reflection_input = {
                "title": paragraph.title,
                "content": paragraph.content,
                "paragraph_latest_state": paragraph.research.latest_summary
            }
            
            # 生成反思搜索查询
            reflection_output = self.reflection_node.run(reflection_input)
            search_query = reflection_output["search_query"]
            search_tool = reflection_output.get("search_tool", "basic_search_news")  # 默认工具
            reasoning = reflection_output["reasoning"]
            
            logger.info(f"    反思查询: {search_query}")
            logger.info(f"    选择的工具: {search_tool}")
            logger.info(f"    反思推理: {reasoning}")
            
            # 执行反思搜索
            # 处理search_news_by_date的特殊参数
            search_kwargs = {}
            if search_tool == "search_news_by_date":
                start_date = reflection_output.get("start_date")
                end_date = reflection_output.get("end_date")
                
                if start_date and end_date:
                    # 验证日期格式
                    if self._validate_date_format(start_date) and self._validate_date_format(end_date):
                        search_kwargs["start_date"] = start_date
                        search_kwargs["end_date"] = end_date
                        logger.info(f"    时间范围: {start_date} 到 {end_date}")
                    else:
                        logger.info(f"    ⚠️  日期格式错误（应为YYYY-MM-DD），改用基础搜索")
                        logger.info(f"        提供的日期: start_date={start_date}, end_date={end_date}")
                        search_tool = "basic_search_news"
                else:
                    logger.info(f"    ⚠️  search_news_by_date工具缺少时间参数，改用基础搜索")
                    search_tool = "basic_search_news"
            
            search_response = self.execute_search_tool(search_tool, search_query, **search_kwargs)
            
            # 转换为兼容格式
            search_results = []
            if search_response and search_response.results:
                configured_cap = int(getattr(self.config, "MAX_SEARCH_RESULTS", 10) or 10)
                configured_cap = max(1, configured_cap)
                max_results = min(len(search_response.results), configured_cap)
                for result in search_response.results[:max_results]:
                    search_results.append({
                        'title': result.title,
                        'url': result.url,
                        'content': result.content,
                        'score': result.score,
                        'raw_content': result.raw_content,
                        'published_date': result.published_date
                    })
            
            if search_results:
                logger.info(f"    找到 {len(search_results)} 个反思搜索结果")
                for j, result in enumerate(search_results, 1):
                    date_info = f" (发布于: {result.get('published_date', 'N/A')})" if result.get('published_date') else ""
                    logger.info(f"      {j}. {result['title'][:50]}...{date_info}")
            else:
                logger.info("    未找到反思搜索结果")
            
            # 更新搜索历史
            paragraph.research.add_search_results(search_query, search_results)
            
            # 生成反思总结
            reflection_summary_input = {
                "title": paragraph.title,
                "content": paragraph.content,
                "search_query": search_query,
                "search_results": format_search_results_for_prompt(
                    search_results, self.config.SEARCH_CONTENT_MAX_LENGTH
                ),
                "paragraph_latest_state": paragraph.research.latest_summary
            }
            
            # 更新状态
            self.state = self.reflection_summary_node.mutate_state(
                reflection_summary_input, self.state, paragraph_index
            )
            
            logger.info(f"    反思 {reflection_i + 1} 完成")
    
    def _generate_final_report(self) -> str:
        """生成最终报告"""
        logger.info(f"\n[步骤 3] 生成最终报告...")
        
        # 准备报告数据
        report_data = []
        for paragraph in self.state.paragraphs:
            report_data.append({
                "title": paragraph.title,
                "paragraph_latest_state": paragraph.research.latest_summary
            })
        
        # 格式化报告
        try:
            final_report = self.report_formatting_node.run(report_data)
        except Exception as e:
            logger.error(f"LLM格式化失败，使用备用方法: {str(e)}")
            final_report = self.report_formatting_node.format_report_manually(
                report_data, self.state.report_title
            )
        
        # 更新状态
        self.state.final_report = final_report
        self.state.mark_completed()
        
        logger.info("最终报告生成完成")
        return final_report
    
    def _save_report(self, report_content: str, export_formats: List[str] = None):
        """
        保存报告到文件，支持多种格式导出
        每个任务的报告保存到独立文件夹中
        
        Args:
            report_content: 报告Markdown内容
            export_formats: 要导出的格式列表，支持 ['md', 'html', 'pdf', 'docx']
        """
        import uuid
        
        # 生成任务ID和时间戳
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        task_id = uuid.uuid4().hex[:8]
        query_safe = "".join(c for c in self.state.query if c.isalnum() or c in (' ', '-', '_')).rstrip()
        query_safe = query_safe.replace(' ', '_')[:30]
        
        # 创建任务专属文件夹
        task_folder_name = f"task_{query_safe}_{timestamp}_{task_id}"
        task_folder = os.path.join(self.config.OUTPUT_DIR, task_folder_name)
        os.makedirs(task_folder, exist_ok=True)
        
        # 默认导出格式
        if export_formats is None:
            export_formats = ['md', 'html']  # 默认导出Markdown和HTML
        
        saved_files = {}
        
        # 保存Markdown
        if 'md' in export_formats:
            md_filepath = os.path.join(task_folder, "report.md")
            with open(md_filepath, 'w', encoding='utf-8') as f:
                f.write(report_content)
            saved_files['md'] = md_filepath
            logger.info(f"Markdown报告已保存到: {md_filepath}")
        
        # 导出HTML
        if 'html' in export_formats:
            try:
                html_content = self._convert_to_html(report_content)
                html_filepath = os.path.join(task_folder, "report.html")
                with open(html_filepath, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                saved_files['html'] = html_filepath
                logger.info(f"HTML报告已保存到: {html_filepath}")
            except Exception as e:
                logger.error(f"HTML导出失败: {e}")
        
        # 导出PDF
        if 'pdf' in export_formats:
            try:
                pdf_filepath = os.path.join(task_folder, "report.pdf")
                self._convert_to_pdf(report_content, pdf_filepath)
                saved_files['pdf'] = pdf_filepath
                logger.info(f"PDF报告已保存到: {pdf_filepath}")
            except Exception as e:
                logger.error(f"PDF导出失败: {e}")
        
        # 导出DOCX
        if 'docx' in export_formats:
            try:
                docx_filepath = os.path.join(task_folder, "report.docx")
                self._convert_to_docx(report_content, docx_filepath)
                saved_files['docx'] = docx_filepath
                logger.info(f"Word报告已保存到: {docx_filepath}")
            except Exception as e:
                logger.error(f"Word导出失败: {e}")
        
        # 保存状态（如果配置允许）
        if self.config.SAVE_INTERMEDIATE_STATES:
            state_filepath = os.path.join(task_folder, "state.json")
            self.state.save_to_file(state_filepath)
            logger.info(f"状态已保存到: {state_filepath}")
        
        # 保存任务清单manifest
        manifest = {
            "task_id": task_id,
            "query": self.state.query,
            "timestamp": timestamp,
            "files": saved_files,
            "folder": task_folder
        }
        manifest_path = os.path.join(task_folder, "manifest.json")
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        
        logger.info(f"任务报告已保存到文件夹: {task_folder}")
        return saved_files
    
    def _convert_to_html(self, markdown_content: str) -> str:
        """将Markdown转换为HTML"""
        try:
            import markdown
            from markdown.extensions.tables import TableExtension
            from markdown.extensions.fenced_code import FencedCodeExtension
            from markdown.extensions.toc import TocExtension
        except ImportError:
            logger.warning("markdown库未安装，使用简单HTML转换")
            return self._simple_markdown_to_html(markdown_content)
        
        # 创建Markdown转换器
        md = markdown.Markdown(extensions=[
            'tables',
            'fenced_code',
            'toc',
            'nl2br',
            'sane_lists',
        ])
        
        html_body = md.convert(markdown_content)
        
        # 包装完整HTML
        html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>深度搜索报告</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.8;
            color: #333;
            background-color: #fafafa;
        }}
        h1 {{ color: #1a1a1a; border-bottom: 3px solid #2196F3; padding-bottom: 15px; }}
        h2 {{ color: #333; border-bottom: 2px solid #eee; padding-bottom: 10px; margin-top: 30px; }}
        h3 {{ color: #444; margin-top: 25px; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            background: white;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px 15px;
            text-align: left;
        }}
        th {{ background-color: #2196F3; color: white; }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Consolas', 'Monaco', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #2196F3;
            margin: 20px 0;
            padding: 10px 20px;
            background-color: #f9f9f9;
            color: #555;
        }}
        strong {{ color: #1a1a1a; }}
        a {{ color: #2196F3; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        ul, ol {{ padding-left: 25px; }}
        li {{ margin: 8px 0; }}
        .report-meta {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 30px;
        }}
        @media print {{
            body {{ background: white; max-width: none; }}
            .no-print {{ display: none; }}
        }}
    </style>
</head>
<body>
{html_body}
<footer style="margin-top: 50px; padding-top: 20px; border-top: 1px solid #eee; color: #888; font-size: 12px;">
    <p>本报告由深度搜索系统自动生成 | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
</footer>
</body>
</html>"""
        
        return html_template
    
    def _simple_markdown_to_html(self, content: str) -> str:
        """简单的Markdown到HTML转换（不依赖外部库）"""
        import html as html_module
        
        # 转义HTML
        content = html_module.escape(content)
        
        # 标题转换
        content = re.sub(r'^### (.+)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^# (.+)$', r'<h1>\1</h1>', content, flags=re.MULTILINE)
        
        # 粗体和斜体
        content = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', content)
        content = re.sub(r'\*(.+?)\*', r'<em>\1</em>', content)
        
        # 链接
        content = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', content)
        
        # 段落
        content = re.sub(r'\n\n', '</p><p>', content)
        content = f'<p>{content}</p>'
        
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>深度搜索报告</title>
    <style>
        body {{ font-family: sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        a {{ color: #0066cc; }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
    
    def _convert_to_pdf(self, markdown_content: str, output_path: str):
        """将Markdown转换为PDF"""
        try:
            from weasyprint import HTML
            
            # 先转换为HTML
            html_content = self._convert_to_html(markdown_content)
            
            # 转换为PDF
            HTML(string=html_content).write_pdf(output_path)
            logger.info(f"PDF生成成功: {output_path}")
            
        except ImportError:
            logger.error("weasyprint库未安装，无法生成PDF。请运行: pip install weasyprint")
            raise ImportError("weasyprint未安装")
        except Exception as e:
            logger.error(f"PDF生成失败: {e}")
            raise
    
    def _convert_to_docx(self, markdown_content: str, output_path: str):
        """将Markdown转换为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.error("python-docx库未安装，无法生成Word文档。请运行: pip install python-docx")
            raise ImportError("python-docx未安装")
        
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        
        # 解析Markdown并转换
        lines = markdown_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                current_paragraph = None
                continue
            
            # 处理标题
            if stripped.startswith('# '):
                heading = doc.add_heading(stripped[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            # 处理列表
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', stripped):
                content = re.sub(r'^\d+\. ', '', stripped)
                doc.add_paragraph(content, style='List Number')
            # 处理引用
            elif stripped.startswith('> '):
                para = doc.add_paragraph(stripped[2:])
                para.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
            # 处理普通段落
            else:
                # 处理粗体
                if '**' in stripped:
                    para = doc.add_paragraph()
                    parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = para.add_run(part[2:-2])
                            run.bold = True
                        else:
                            para.add_run(part)
                else:
                    doc.add_paragraph(stripped)
        
        # 保存文档
        doc.save(output_path)
        logger.info(f"Word文档生成成功: {output_path}")
    
    def get_progress_summary(self) -> Dict[str, Any]:
        """获取进度摘要"""
        return self.state.get_progress_summary()
    
    def load_state(self, filepath: str):
        """从文件加载状态"""
        self.state = State.load_from_file(filepath)
        logger.info(f"状态已从 {filepath} 加载")
    
    def save_state(self, filepath: str):
        """保存状态到文件"""
        self.state.save_to_file(filepath)
        logger.info(f"状态已保存到 {filepath}")


def create_agent() -> DeepSearchAgent:
    """
    创建Deep Search Agent实例的便捷函数
    
    Returns:
        DeepSearchAgent实例
    """
    from .utils.config import Settings
    config = Settings()
    return DeepSearchAgent(config)
